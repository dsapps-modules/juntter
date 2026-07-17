from pathlib import Path
import math
import textwrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


def u(value: str) -> str:
    return value.encode("ascii").decode("unicode_escape")


BASE = Path(r"C:\Users\User\Desktop\juntter")
ASSETS = BASE / "docs_unpacked" / "docs"
OUT_DIR = Path(r"C:\xampp\htdocs\juntter\docs")
OUT_DIR.mkdir(parents=True, exist_ok=True)

FONT_REGULAR = r"C:\Windows\Fonts\arial.ttf"
FONT_BOLD = r"C:\Windows\Fonts\arialbd.ttf"

PAGE_W = Mm(210)
PAGE_H = Mm(297)
MARGIN_MM = 17
CONTENT_W = Inches(8.27 - (MARGIN_MM / 25.4) * 2)

BG = "#F7F7F6"
BLACK = "111111"
MID = "666666"
GOLD = "FFC400"
GOLD_DARK = "C79200"
WHITE = "#FFFFFF"


def font(path: str, size: int) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(path, size=size)


def pcolor(color: str) -> str:
    if color.startswith("#"):
        return color
    return f"#{color}"


def add_text(draw: ImageDraw.ImageDraw, xy, text, ft, fill, anchor="la", spacing=8, align="left"):
    draw.multiline_text(
        xy,
        text,
        font=ft,
        fill=pcolor(fill),
        spacing=spacing,
        anchor=anchor,
        align=align,
    )


def rounded_rect(draw, box, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def set_run_font(run, name="Arial", size=11, bold=False, color=BLACK, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ["ascii", "hAnsi", "eastAsia", "cs"]:
        rfonts.set(qn(f"w:{attr}"), name)


def set_spacing(paragraph, before=0, after=0, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def style_doc(doc: Document) -> None:
    styles = doc.styles
    for name in ["Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
        try:
            style = styles[name]
            style.font.name = "Arial"
            style.font.color.rgb = RGBColor.from_string(BLACK)
        except Exception:
            pass

    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(5)
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.bold = True
    styles["Heading 3"].font.size = Pt(11)
    styles["Heading 3"].font.bold = True


def add_header_footer(section) -> None:
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo = header.add_run()
    logo.add_picture(str(ASSETS / "logo" / "juntter_png_256.png"), width=Inches(0.45))
    text = header.add_run(u("  Manual do Administrador | Checkout Juntter"))
    set_run_font(text, size=9, bold=True, color=BLACK)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_text = footer.add_run(u("Checkout Juntter - Manual do Administrador  "))
    set_run_font(footer_text, size=8.5, color=MID)
    add_page_number(footer)


def create_cover():
    source = Image.open(ASSETS / "homepage-screenshot.png").convert("RGB")
    hero = source.crop((0, 0, source.width, 1500))

    canvas = Image.new("RGB", (1654, 2339), BG)
    draw = ImageDraw.Draw(canvas)
    rounded_rect(draw, (115, 90, 1539, 1145), 34, WHITE, "#EFEFEF", 2)

    hero = hero.resize((1424, 920))
    canvas.paste(hero, (115, 90))

    ft_logo = font(FONT_BOLD, 34)
    ft_kicker = font(FONT_BOLD, 28)
    ft_title = font(FONT_BOLD, 72)
    ft_sub = font(FONT_REGULAR, 34)
    ft_small = font(FONT_REGULAR, 24)
    ft_chip = font(FONT_BOLD, 23)

    logo = Image.open(ASSETS / "logo" / "juntter_png_256.png").convert("RGBA")
    logo = logo.resize((170, 170))
    canvas.paste(logo, (118, 36), logo)
    add_text(draw, (300, 56), u("Checkout Juntter"), ft_logo, BLACK)

    rounded_rect(draw, (110, 1125, 1544, 2210), 34, WHITE, "#EFEFEF", 2)
    draw.rectangle((150, 1176, 308, 1182), fill=f"#{GOLD}")
    add_text(draw, (150, 1215), u("MANUAL DO ADMINISTRADOR"), ft_kicker, f"#{GOLD}")
    add_text(draw, (150, 1272), u("Checkout Juntter"), ft_title, BLACK)
    subtitle = textwrap.fill(
        u(
            "Opera\\u00e7\\u00e3o de gest\\u00e3o para acompanhar o painel geral, "
            "administrar estabelecimentos, abrir acessos de vendedores, revisar regras de split "
            "e manter a base operacional sob controle."
        ),
        width=44,
    )
    add_text(draw, (150, 1388), subtitle, ft_sub, MID)

    chips = [
        (u("Painel"), u("Vis\\u00e3o consolidada")),
        (u("Cadastros"), u("Estabelecimentos e acessos")),
        (u("Controle"), u("Split, exporta\\u00e7\\u00e3o e seguran\\u00e7a")),
    ]
    x = 150
    y = 1548
    for head, body in chips:
        rounded_rect(draw, (x, y, x + 460, y + 155), 24, "#FAFAFA", "#E7E7E7", 2)
        rounded_rect(draw, (x + 22, y + 22, x + 84, y + 84), 18, f"#{GOLD}")
        add_text(draw, (x + 28, y + 31), head[:1], ft_chip, BLACK)
        add_text(draw, (x + 105, y + 28), head, ft_chip, BLACK)
        add_text(draw, (x + 105, y + 68), body, ft_small, MID)
        x += 486

    draw.line((150, 1745, 1504, 1745), fill="#E6E6E6", width=3)
    add_text(draw, (150, 1785), u("Vers\\u00e3o de refer\\u00eancia: julho de 2026"), ft_small, MID)
    add_text(
        draw,
        (150, 1830),
        u("Este manual organiza as rotinas que o administrador precisa conhecer para operar o sistema com crit\\u00e9rio e consist\\u00eancia."),
        ft_small,
        BLACK,
    )

    canvas.save(OUT_DIR / "manual-admin-cover.png", quality=95)


def create_dashboard_mockup(outfile: Path):
    canvas = Image.new("RGB", (1800, 1320), BG)
    draw = ImageDraw.Draw(canvas)
    ft_title = font(FONT_BOLD, 42)
    ft_sub = font(FONT_REGULAR, 25)
    ft_card = font(FONT_BOLD, 25)
    ft_text = font(FONT_REGULAR, 20)
    ft_small = font(FONT_REGULAR, 18)

    add_text(draw, (80, 60), u("Painel administrativo"), ft_title, BLACK)
    add_text(draw, (80, 120), u("Resumo consolidado por m\\u00eas e ano, com sincroniza\\u00e7\\u00e3o de base em segundo plano."), ft_sub, MID)

    rounded_rect(draw, (70, 200, 1730, 1220), 28, WHITE, "#E6E6E6", 2)
    rounded_rect(draw, (105, 235, 335, 325), 18, "#FFF8D8", f"#{GOLD}", 2)
    add_text(draw, (130, 255), u("Filtro"), ft_card, BLACK)
    add_text(draw, (130, 292), u("Per\\u00edodo: julho de 2026"), ft_text, MID)

    metric_specs = [
        (105, 360, "Transa\\u00e7\\u00f5es", "3.842", "18% acima do m\\u00eas anterior"),
        (485, 360, "Volume bruto", "R$ 1.240.980", "Cart\\u00e3o, PIX e boleto"),
        (865, 360, "Taxas", "R$ 74.421", "M\\u00e9dia ponderada"),
        (1245, 360, "Boletos", "1.026", "Em aberto e pagos"),
    ]
    for x, y, title, value, desc in metric_specs:
        rounded_rect(draw, (x, y, x + 300, y + 145), 20, "#FBFBFB", "#E6E6E6", 2)
        add_text(draw, (x + 22, y + 20), u(title), ft_small, MID)
        add_text(draw, (x + 22, y + 55), value, font(FONT_BOLD, 32), BLACK)
        add_text(draw, (x + 22, y + 105), u(desc), ft_small, MID)

    rounded_rect(draw, (105, 545, 1030, 1095), 20, "#FCFCFC", "#E6E6E6", 2)
    add_text(draw, (130, 570), u("Distribui\\u00e7\\u00e3o por m\\u00e9todo"), ft_card, BLACK)
    for i, h in enumerate([220, 310, 180, 360, 260, 390]):
        x = 150 + i * 120
        draw.rectangle((x, 940 - h, x + 70, 940), fill=f"#{GOLD}")
        add_text(draw, (x, 955), str(i + 1), ft_small, MID)
    draw.line((145, 940, 990, 940), fill="#D8D8D8", width=2)
    for y in [820, 700, 580]:
        draw.line((145, y, 990, y), fill="#F0F0F0", width=1)

    rounded_rect(draw, (1060, 545, 1705, 1095), 20, "#FCFCFC", "#E6E6E6", 2)
    add_text(draw, (1085, 570), u("Status consolidados"), ft_card, BLACK)
    status_lines = [
        (u("PAID"), 48),
        (u("PENDING"), 18),
        (u("PROCESSING"), 12),
        (u("FAILED"), 8),
        (u("REFUNDED"), 4),
    ]
    yy = 635
    for label, pct in status_lines:
        draw.rectangle((1085, yy, 1085 + pct * 4, yy + 24), fill=f"#{GOLD}")
        add_text(draw, (1085, yy - 5), f"{label}", ft_small, BLACK)
        add_text(draw, (1500, yy - 5), f"{pct}%", ft_small, MID)
        yy += 72

    rounded_rect(draw, (105, 1120, 1705, 1190), 18, "#FFF8D8", None, 0)
    add_text(draw, (130, 1143), u("A abertura do painel dispara a sincroniza\\u00e7\\u00e3o Paytime em background e consolida transa\\u00e7\\u00f5es, boletos e saldos por estabelecimento."), ft_text, BLACK)

    canvas.save(outfile, quality=95)


def create_establishments_mockup(outfile: Path):
    canvas = Image.new("RGB", (1800, 1320), BG)
    draw = ImageDraw.Draw(canvas)
    ft_title = font(FONT_BOLD, 42)
    ft_sub = font(FONT_REGULAR, 25)
    ft_head = font(FONT_BOLD, 24)
    ft_small = font(FONT_REGULAR, 18)
    ft_text = font(FONT_REGULAR, 20)

    add_text(draw, (80, 60), u("Estabelecimentos"), ft_title, BLACK)
    add_text(draw, (80, 120), u("Busca, exporta\\u00e7\\u00e3o, consulta de detalhes e edi\\u00e7\\u00e3o de dados cadastrais."), ft_sub, MID)
    rounded_rect(draw, (70, 200, 1730, 1220), 28, WHITE, "#E6E6E6", 2)
    rounded_rect(draw, (105, 235, 520, 1170), 20, "#FCFCFC", "#E6E6E6", 2)
    add_text(draw, (130, 260), u("Busca"), ft_head, BLACK)
    rounded_rect(draw, (130, 315, 495, 365), 10, "#FFFFFF", "#DADADA", 1)
    add_text(draw, (150, 330), u("Buscar por nome, documento ou e-mail"), ft_small, MID)

    rows = [
        ("Loja Aurora", "active", "R$ 43.200"),
        ("Mercado 7", "review", "R$ 18.930"),
        ("Mega Digital", "inactive", "R$ 2.100"),
        ("Studio Alpha", "active", "R$ 76.500"),
    ]
    y = 400
    for name, status, revenue in rows:
        rounded_rect(draw, (130, y, 495, y + 105), 14, "#FFFFFF", "#EAEAEA", 1)
        add_text(draw, (150, y + 18), name, ft_head, BLACK)
        add_text(draw, (150, y + 55), revenue, ft_text, MID)
        tag_fill = "#D5F3D6" if status == "active" else "#FFF4C4" if status == "review" else "#F0F0F0"
        tag_text = u("Ativo") if status == "active" else u("Em revis\\u00e3o") if status == "review" else u("Inativo")
        rounded_rect(draw, (360, y + 28, 475, y + 72), 14, tag_fill, None, 0)
        add_text(draw, (375, y + 39), tag_text, ft_small, BLACK)
        y += 124

    rounded_rect(draw, (575, 235, 1695, 1170), 20, "#FCFCFC", "#E6E6E6", 2)
    add_text(draw, (600, 260), u("Detalhe selecionado"), ft_head, BLACK)
    rounded_rect(draw, (600, 315, 1085, 420), 16, "#FFF9DD", f"#{GOLD}", 2)
    add_text(draw, (625, 338), u("Tipo de acesso"), ft_small, MID)
    add_text(draw, (625, 366), u("ACQUIRER"), font(FONT_BOLD, 30), BLACK)
    add_text(draw, (840, 338), u("Status"), ft_small, MID)
    add_text(draw, (840, 366), u("Ativo"), font(FONT_BOLD, 30), BLACK)
    add_text(draw, (1020, 338), u("Risco"), ft_small, MID)
    add_text(draw, (1020, 366), u("Baixo"), font(FONT_BOLD, 30), BLACK)

    labels = [
        (610, 460, u("Raz\\u00e3o social / fantasia"), u("Loja Aurora Ltda.")),
        (610, 555, u("Documento"), u("12.345.678/0001-90")),
        (610, 650, u("E-mail"), u("contato@lojaaurora.com")),
        (610, 745, u("Telefone"), u("(11) 99999-0000")),
        (610, 840, u("Receita"), u("R$ 43.200,00")),
        (610, 935, u("Formato"), u("LTDA")),
    ]
    for x, y0, label, value in labels:
        add_text(draw, (x, y0), label, ft_small, MID)
        rounded_rect(draw, (x, y0 + 30, x + 440, y0 + 75), 12, WHITE, "#DADADA", 1)
        add_text(draw, (x + 16, y0 + 42), value, ft_text, BLACK)

    rounded_rect(draw, (1100, 460, 1660, 860), 16, "#FFFFFF", "#EAEAEA", 1)
    add_text(draw, (1125, 490), u("A\\u00e7\\u00f5es r\\u00e1pidas"), ft_head, BLACK)
    for idx, label in enumerate([u("Editar dados"), u("Exportar XLS"), u("Criar split"), u("Consultar split")]):
        yy = 545 + idx * 72
        rounded_rect(draw, (1125, yy, 1390, yy + 46), 12, "#FFF8D8", f"#{GOLD}", 1)
        add_text(draw, (1142, yy + 12), label, ft_small, BLACK)

    rounded_rect(draw, (1100, 900, 1660, 1110), 16, "#FFF8D8", f"#{GOLD}", 1)
    add_text(draw, (1125, 930), u("Exporta\\u00e7\\u00e3o"), ft_head, BLACK)
    add_text(draw, (1125, 980), u("Gera planilha .xls com a base filtrada e ordenada para an\\u00e1lise externa."), ft_text, BLACK)

    canvas.save(outfile, quality=95)


def create_split_mockup(outfile: Path):
    canvas = Image.new("RGB", (1800, 1320), BG)
    draw = ImageDraw.Draw(canvas)
    ft_title = font(FONT_BOLD, 42)
    ft_sub = font(FONT_REGULAR, 25)
    ft_head = font(FONT_BOLD, 24)
    ft_text = font(FONT_REGULAR, 20)
    ft_small = font(FONT_REGULAR, 18)

    add_text(draw, (80, 60), u("Split pr\\u00e9"), ft_title, BLACK)
    add_text(draw, (80, 120), u("Cria\\u00e7\\u00e3o, consulta e remo\\u00e7\\u00e3o de regras de rateio por estabelecimento."), ft_sub, MID)
    rounded_rect(draw, (70, 200, 1730, 1220), 28, WHITE, "#E6E6E6", 2)

    rounded_rect(draw, (105, 235, 980, 1160), 20, "#FCFCFC", "#E6E6E6", 2)
    add_text(draw, (130, 260), u("Nova regra"), ft_head, BLACK)
    form_fields = [
        (u("T\\u00edtulo"), u("Repasse principal")),
        (u("Modalidade"), u("ALL / CREDIT / DEBIT / PIX")),
        (u("Canal"), u("ALL / CHIP / TAP / SMART / ONLINE")),
        (u("Divis\\u00e3o"), u("PERCENTAGE ou CURRENCY")),
        (u("Parcelas"), u("1 a 12")),
    ]
    yy = 320
    for label, value in form_fields:
        add_text(draw, (130, yy), label, ft_small, MID)
        rounded_rect(draw, (130, yy + 28, 920, yy + 76), 12, WHITE, "#DADADA", 1)
        add_text(draw, (150, yy + 41), value, ft_text, BLACK)
        yy += 105
    rounded_rect(draw, (130, 900, 920, 1095), 16, "#FFF8D8", f"#{GOLD}", 1)
    add_text(draw, (155, 930), u("Use a lista de estabelecimentos para distribuir o valor. Cada item precisa de id, valor e status de ativo."), ft_text, BLACK)

    rounded_rect(draw, (1035, 235, 1670, 1160), 20, "#FCFCFC", "#E6E6E6", 2)
    add_text(draw, (1060, 260), u("Benefici\\u00e1rios"), ft_head, BLACK)
    headers = [u("Estabelecimento"), u("Valor"), u("Ativo")]
    x_positions = [1060, 1355, 1520]
    for x, h in zip(x_positions, headers):
        add_text(draw, (x, 315), h, ft_small, MID)
    table_rows = [
        ("Loja Aurora", "60%", "Sim"),
        ("Studio Alpha", "25%", "Sim"),
        ("Parceiro B2B", "15%", "N\\u00e3o"),
    ]
    y = 355
    for row in table_rows:
        rounded_rect(draw, (1060, y, 1640, y + 120), 12, WHITE, "#EAEAEA", 1)
        add_text(draw, (1080, y + 20), row[0], ft_text, BLACK)
        add_text(draw, (1380, y + 20), row[1], ft_text, BLACK)
        add_text(draw, (1545, y + 20), row[2], ft_text, BLACK)
        y += 140
    rounded_rect(draw, (1060, 880, 1640, 1088), 16, "#FFF8D8", f"#{GOLD}", 1)
    add_text(draw, (1085, 910), u("Cuidado"), ft_head, BLACK)
    add_text(draw, (1085, 965), u("Alterar uma regra de split impacta o destino financeiro das transa\\u00e7\\u00f5es. Confirme o cen\\u00e1rio antes de gravar."), ft_text, BLACK)

    canvas.save(outfile, quality=95)


def create_vendors_mockup(outfile: Path):
    canvas = Image.new("RGB", (1800, 1320), BG)
    draw = ImageDraw.Draw(canvas)
    ft_title = font(FONT_BOLD, 42)
    ft_sub = font(FONT_REGULAR, 25)
    ft_head = font(FONT_BOLD, 24)
    ft_text = font(FONT_REGULAR, 20)
    ft_small = font(FONT_REGULAR, 18)

    add_text(draw, (80, 60), u("Vendedores"), ft_title, BLACK)
    add_text(draw, (80, 120), u("Cadastro de acessos, atualiza\\u00e7\\u00e3o de dados, senha e vincula\\u00e7\\u00e3o ao estabelecimento."), ft_sub, MID)
    rounded_rect(draw, (70, 200, 1730, 1220), 28, WHITE, "#E6E6E6", 2)

    rounded_rect(draw, (105, 235, 560, 1160), 20, "#FCFCFC", "#E6E6E6", 2)
    add_text(draw, (130, 260), u("Lista de acessos"), ft_head, BLACK)
    sellers = [
        ("Ana Costa", "Loja Aurora", "Admin de loja"),
        ("Bruno Dias", "Studio Alpha", "Admin de loja"),
        ("Carla Lima", "Mercado 7", "Vendedor"),
    ]
    y = 330
    for name, shop, role in sellers:
        rounded_rect(draw, (130, y, 530, y + 115), 14, WHITE, "#EAEAEA", 1)
        add_text(draw, (150, y + 20), name, ft_head, BLACK)
        add_text(draw, (150, y + 58), shop, ft_text, MID)
        rounded_rect(draw, (360, y + 24, 510, y + 68), 14, "#FFF8D8", f"#{GOLD}", 1)
        add_text(draw, (375, y + 36), role, ft_small, BLACK)
        y += 140

    rounded_rect(draw, (610, 235, 1670, 830), 20, "#FCFCFC", "#E6E6E6", 2)
    add_text(draw, (635, 260), u("Novo acesso"), ft_head, BLACK)
    fields = [
        (u("Nome"), u("Nome do respons\\u00e1vel")),
        (u("E-mail"), u("usuario@lojax.com")),
        (u("Senha"), u("Senha inicial")),
        (u("Estabelecimento"), u("Selecionar estabelecimento dispon\\u00edvel")),
    ]
    yy = 320
    for label, value in fields:
        add_text(draw, (635, yy), label, ft_small, MID)
        rounded_rect(draw, (635, yy + 28, 1605, yy + 76), 12, WHITE, "#DADADA", 1)
        add_text(draw, (655, yy + 41), value, ft_text, BLACK)
        yy += 95
    rounded_rect(draw, (635, 705, 980, 780), 14, "#FFF8D8", f"#{GOLD}", 1)
    add_text(draw, (655, 728), u("Criar acesso"), ft_head, BLACK)
    rounded_rect(draw, (1010, 705, 1605, 780), 14, "#F6F6F6", "#DADADA", 1)
    add_text(draw, (1030, 728), u("A redefinir senha e remover acesso"), ft_head, BLACK)

    rounded_rect(draw, (610, 860, 1670, 1160), 20, "#FFF9DD", f"#{GOLD}", 1)
    add_text(draw, (635, 888), u("Fluxo de controle"), ft_head, BLACK)
    add_text(
        draw,
        (635, 940),
        u("O sistema cria o usu\\u00e1rio como vendedor e vincula o acesso ao estabelecimento escolhido. A busca retorna apenas lojas ainda n\\u00e3o usadas por outro acesso."),
        ft_text,
        BLACK,
    )
    canvas.save(outfile, quality=95)


def create_faturamento_mockup(outfile: Path):
    canvas = Image.new("RGB", (1800, 1320), BG)
    draw = ImageDraw.Draw(canvas)
    ft_title = font(FONT_BOLD, 42)
    ft_sub = font(FONT_REGULAR, 25)
    ft_head = font(FONT_BOLD, 24)
    ft_text = font(FONT_REGULAR, 20)
    ft_small = font(FONT_REGULAR, 18)

    add_text(draw, (80, 60), u("Faturamento por vendedor"), ft_title, BLACK)
    add_text(draw, (80, 120), u("Comparativo por loja para identificar volume, ticket m\\u00e9dio, status e oportunidade de revis\\u00e3o."), ft_sub, MID)
    rounded_rect(draw, (70, 200, 1730, 1220), 28, WHITE, "#E6E6E6", 2)

    cards = [
        (105, 250, u("Loja Aurora"), u("R$ 120.450"), u("482 vendas")),
        (475, 250, u("Studio Alpha"), u("R$ 86.240"), u("311 vendas")),
        (845, 250, u("Mercado 7"), u("R$ 28.610"), u("105 vendas")),
        (1215, 250, u("Total"), u("R$ 235.300"), u("898 vendas")),
    ]
    for x, y, title, value, desc in cards:
        rounded_rect(draw, (x, y, x + 300, y + 145), 18, "#FBFBFB", "#E6E6E6", 2)
        add_text(draw, (x + 20, y + 20), title, ft_small, MID)
        add_text(draw, (x + 20, y + 55), value, font(FONT_BOLD, 30), BLACK)
        add_text(draw, (x + 20, y + 105), desc, ft_small, MID)

    rounded_rect(draw, (105, 440, 1180, 1110), 20, "#FCFCFC", "#E6E6E6", 2)
    add_text(draw, (130, 465), u("Tend\\u00eancia mensal"), ft_head, BLACK)
    for i, h in enumerate([180, 220, 160, 260, 290, 340]):
        x = 165 + i * 140
        draw.rectangle((x, 900 - h, x + 90, 900), fill=f"#{GOLD}")
        add_text(draw, (x, 915), f"{i+1}", ft_small, MID)
    draw.line((150, 900, 1130, 900), fill="#D8D8D8", width=2)

    rounded_rect(draw, (1215, 440, 1660, 1110), 20, "#FCFCFC", "#E6E6E6", 2)
    add_text(draw, (1240, 465), u("Alertas"), ft_head, BLACK)
    alert_lines = [
        u("Loja sem movimento h\\u00e1 12 dias."),
        u("3 acessos sem senha trocada."),
        u("2 estabelecimentos com split pendente."),
        u("1 cadastro com e-mail sem confirma\\u00e7\\u00e3o."),
    ]
    y = 545
    for line in alert_lines:
        rounded_rect(draw, (1240, y, 1610, y + 95), 14, "#FFF8D8", f"#{GOLD}", 1)
        add_text(draw, (1260, y + 20), line, ft_text, BLACK)
        y += 120

    canvas.save(outfile, quality=95)


def create_profile_mockup(outfile: Path):
    canvas = Image.new("RGB", (1800, 1320), BG)
    draw = ImageDraw.Draw(canvas)
    ft_title = font(FONT_BOLD, 42)
    ft_sub = font(FONT_REGULAR, 25)
    ft_head = font(FONT_BOLD, 24)
    ft_text = font(FONT_REGULAR, 20)
    ft_small = font(FONT_REGULAR, 18)

    add_text(draw, (80, 60), u("Perfil e seguran\\u00e7a"), ft_title, BLACK)
    add_text(draw, (80, 120), u("Dados do administrador, senha, verifica\\u00e7\\u00e3o de e-mail e identidade visual da conta."), ft_sub, MID)
    rounded_rect(draw, (70, 200, 1730, 1220), 28, WHITE, "#E6E6E6", 2)

    rounded_rect(draw, (105, 235, 1030, 1160), 20, "#FCFCFC", "#E6E6E6", 2)
    add_text(draw, (130, 260), u("Dados da conta"), ft_head, BLACK)
    fields = [
        (u("Nome"), u("Administrador Juntter")),
        (u("E-mail"), u("admin@juntter.com")),
        (u("Status"), u("Verificado")),
        (u("Senha"), u("Alterar senha")),
    ]
    yy = 320
    for label, value in fields:
        add_text(draw, (130, yy), label, ft_small, MID)
        rounded_rect(draw, (130, yy + 28, 980, yy + 76), 12, WHITE, "#DADADA", 1)
        add_text(draw, (150, yy + 41), value, ft_text, BLACK)
        yy += 105

    rounded_rect(draw, (130, 770, 980, 1090), 16, "#FFF8D8", f"#{GOLD}", 1)
    add_text(draw, (155, 800), u("Logotipo"), ft_head, BLACK)
    add_text(draw, (155, 850), u("Carregar ou remover a imagem da empresa usada nos cadastros e nas telas p\\u00fablicas."), ft_text, BLACK)

    rounded_rect(draw, (1070, 235, 1665, 1160), 20, "#FCFCFC", "#E6E6E6", 2)
    add_text(draw, (1095, 260), u("Boas pr\\u00e1ticas"), ft_head, BLACK)
    tips = [
        u("Troque a senha periodicamente."),
        u("Revise o e-mail sempre que houver altera\\u00e7\\u00e3o de contato."),
        u("Reenvie a verifica\\u00e7\\u00e3o quando o acesso ficar pendente."),
        u("Mantenha o logotipo leg\\u00edvel e no formato adequado para tela."),
    ]
    y = 330
    for tip in tips:
        rounded_rect(draw, (1095, y, 1635, y + 95), 14, "#FFF9DD", f"#{GOLD}", 1)
        add_text(draw, (1115, y + 20), tip, ft_text, BLACK)
        y += 120

    canvas.save(outfile, quality=95)


def add_cover_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(str(OUT_DIR / "manual-admin-cover.png"), width=CONTENT_W)


def add_simple_page(doc, kicker, title, paragraphs, bullets=None, note=None):
    p = doc.add_paragraph()
    r = p.add_run(kicker.upper())
    set_run_font(r, size=9, bold=True, color=GOLD)

    p = doc.add_paragraph()
    r = p.add_run(title)
    set_run_font(r, size=22, bold=True, color=BLACK)

    for text in paragraphs:
        p = doc.add_paragraph()
        set_spacing(p, after=6, line=1.25)
        r = p.add_run(text)
        set_run_font(r, size=10.5, color="444444")

    if bullets:
        for text in bullets:
            p = doc.add_paragraph(style="List Bullet")
            set_spacing(p, after=1, line=1.2)
            p.paragraph_format.left_indent = Inches(0.18)
            r = p.add_run(text)
            set_run_font(r, size=10.5, color=BLACK)

    if note:
        p = doc.add_paragraph()
        set_shading(p, "F0F0F0")
        r = p.add_run(note["label"] + ": ")
        set_run_font(r, size=10.5, bold=True, color=BLACK)
        r = p.add_run(note["text"])
        set_run_font(r, size=10.5, color=BLACK)


def add_section_page(doc, kicker, title, intro, image_path=None, bullets=None, steps=None, note=None, steps_style="List Number"):
    p = doc.add_paragraph()
    r = p.add_run(kicker.upper())
    set_run_font(r, size=9, bold=True, color=GOLD)

    p = doc.add_paragraph()
    r = p.add_run(title)
    set_run_font(r, size=22, bold=True, color=BLACK)

    p = doc.add_paragraph()
    set_spacing(p, after=8, line=1.25)
    r = p.add_run(intro)
    set_run_font(r, size=10.5, color="444444")

    if image_path:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(5)
        p.add_run().add_picture(str(image_path), width=CONTENT_W)

    if note:
        p = doc.add_paragraph()
        set_shading(p, "FFF7D1")
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(note["label"] + ": ")
        set_run_font(r, size=10.5, bold=True, color=BLACK)
        r = p.add_run(note["text"])
        set_run_font(r, size=10.5, color=BLACK)

    if steps:
        for text in steps:
            p = doc.add_paragraph(style=steps_style)
            set_spacing(p, after=2, line=1.2)
            p.paragraph_format.left_indent = Inches(0.25)
            r = p.add_run(text)
            set_run_font(r, size=10.5, color=BLACK)

    if bullets:
        for text in bullets:
            p = doc.add_paragraph(style="List Bullet")
            set_spacing(p, after=1, line=1.2)
            p.paragraph_format.left_indent = Inches(0.18)
            r = p.add_run(text)
            set_run_font(r, size=10.5, color=BLACK)


def build_images():
    create_cover()
    create_dashboard_mockup(OUT_DIR / "admin-dashboard.png")
    create_establishments_mockup(OUT_DIR / "admin-establishments.png")
    create_split_mockup(OUT_DIR / "admin-split.png")
    create_vendors_mockup(OUT_DIR / "admin-vendors.png")
    create_faturamento_mockup(OUT_DIR / "admin-faturamento.png")
    create_profile_mockup(OUT_DIR / "admin-profile.png")


def build_docx():
    doc = Document()
    style_doc(doc)

    sec = doc.sections[0]
    sec.page_width = PAGE_W
    sec.page_height = PAGE_H
    sec.left_margin = Mm(MARGIN_MM)
    sec.right_margin = Mm(MARGIN_MM)
    sec.top_margin = Mm(12)
    sec.bottom_margin = Mm(12)
    sec.different_first_page_header_footer = True
    add_header_footer(sec)

    add_cover_page(doc)
    doc.add_page_break()

    add_simple_page(
        doc,
        u("Dire\\u00e7\\u00e3o editorial"),
        u("Como ler este manual"),
        [
            u("Este manual foi escrito para apoiar a rotina do administrador. Cada se\\u00e7\\u00e3o responde ao que acompanhar, quando intervir e como concluir a tarefa."),
            u("A leitura segue a ordem natural do trabalho: painel, estabelecimentos, vendedores, regras de split, perfil e checklist operacional."),
            u("O visual permanece limpo, com preto, branco e dourado, para favorecer leitura r\\u00e1pida e navega\\u00e7\\u00e3o segura."),
        ],
        bullets=[
            u("Use o painel para identificar o estado geral da opera\\u00e7\\u00e3o."),
            u("Use Estabelecimentos para localizar, exportar e editar a base."),
            u("Use Vendedores para abrir, ajustar ou remover acessos."),
            u("Use Split pr\\u00e9 para administrar o repasse entre participantes."),
        ],
        note={
            "label": u("Princ\\u00edpio central"),
            "text": u("se uma configura\\u00e7\\u00e3o altera o fluxo financeiro, revise com cuidado antes de salvar."),
        },
    )

    doc.add_page_break()
    add_section_page(
        doc,
        u("Painel"),
        u("Dashboard administrativo"),
        u("Ao entrar no painel administrativo, o sistema consolida transa\\u00e7\\u00f5es, boletos, taxas, ticket m\\u00e9dio e distribui\\u00e7\\u00e3o por status. A sincroniza\\u00e7\\u00e3o em background pode ser disparada ao abrir a tela."),
        image_path=OUT_DIR / "admin-dashboard.png",
        steps=[
            u("Escolha m\\u00eas e ano antes de comparar per\\u00edodos."),
            u("Confira volume bruto, taxas, boletos e total de transa\\u00e7\\u00f5es."),
            u("Use a distribui\\u00e7\\u00e3o por status para localizar falhas ou pend\\u00eancias."),
            u("Acompanhe o painel como leitura executiva, n\\u00e3o apenas como listagem de dados."),
        ],
        note={
            "label": u("Observa\\u00e7\\u00e3o"),
            "text": u("o dashboard do administrador consolida informa\\u00e7\\u00f5es da base inteira; o super admin usa a mesma l\\u00f3gica com vis\\u00e3o ainda mais ampla."),
        },
    )

    doc.add_page_break()
    add_section_page(
        doc,
        u("Estabelecimentos"),
        u("Busca, exporta\\u00e7\\u00e3o e edi\\u00e7\\u00e3o cadastral"),
        u("A \\u00e1rea de estabelecimentos \\u00e9 usada para localizar uma loja, abrir o detalhe, revisar dados principais e ajustar informa\\u00e7\\u00f5es cadastrais quando houver mudan\\u00e7a operacional."),
        image_path=OUT_DIR / "admin-establishments.png",
        bullets=[
            u("Pesquise por nome fantasia, respons\\u00e1vel, documento ou e-mail."),
            u("Abra o detalhe para ver tipo de acesso, status, risco, receita e endere\\u00e7o."),
            u("Edite access_type, nome, telefone, receita, formato, e-mail, GMV e data de nascimento."),
            u("Exporte a lista em .xls quando precisar de an\\u00e1lise externa."),
        ],
        note={
            "label": u("Cuidado"),
            "text": u("o mesmo cadastro pode impactar outros fluxos do sistema; revise o documento e os dados financeiros antes de confirmar a edi\\u00e7\\u00e3o."),
        },
    )

    doc.add_page_break()
    add_section_page(
        doc,
        u("Split pr\\u00e9"),
        u("Regras de rateio por estabelecimento"),
        u("As regras de split definem como o valor de uma transa\\u00e7\\u00e3o ser\\u00e1 dividido entre participantes. A cria\\u00e7\\u00e3o exige t\\u00edtulo, modalidade, canal, tipo de divis\\u00e3o, parcela e a lista dos benefici\\u00e1rios."),
        image_path=OUT_DIR / "admin-split.png",
        steps_style="List Number 2",
        steps=[
            u("Defina a modalidade: ALL, CREDIT, DEBIT ou PIX."),
            u("Escolha o canal: ALL, CHIP, TAP, SMART ou ONLINE."),
            u("Informe PERCENTAGE ou CURRENCY conforme o acordo comercial."),
            u("Inclua os estabelecimentos com id, valor e ativo."),
            u("Consulte e exclua regras somente quando tiver certeza do impacto no repasse."),
        ],
        note={
            "label": u("Importante"),
            "text": u("uma altera\\u00e7\\u00e3o de split interfere no destino do valor liquidado; confirme antes de salvar ou remover uma regra."),
        },
    )

    doc.add_page_break()
    add_section_page(
        doc,
        u("Vendedores"),
        u("Acessos, senhas e vincula\\u00e7\\u00e3o de loja"),
        u("O administrador cria e mant\\u00e9m os acessos dos vendedores. A busca lista apenas estabelecimentos dispon\\u00edveis, isto \\u00e9, que ainda n\\u00e3o possuem um acesso vendedor associado."),
        image_path=OUT_DIR / "admin-vendors.png",
        steps_style="List Number 3",
        steps=[
            u("Selecione o estabelecimento que ainda n\\u00e3o foi usado em outro acesso."),
            u("Crie o usu\\u00e1rio com nome, e-mail, senha e estabelecimento."),
            u("Atualize nome e e-mail quando o respons\\u00e1vel mudar."),
            u("Redefina a senha quando houver solicita\\u00e7\\u00e3o de suporte ou seguran\\u00e7a."),
            u("Remova o acesso apenas quando ele n\\u00e3o for mais necess\\u00e1rio."),
        ],
        note={
            "label": u("Regra pr\\u00e1tica"),
            "text": u("o sistema cria o perfil como vendedor vinculado ao estabelecimento e com sub-n\\u00edvel de admin de loja."),
        },
    )

    doc.add_page_break()
    add_section_page(
        doc,
        u("Faturamento"),
        u("Leitura de desempenho por loja"),
        u("A vis\\u00e3o de faturamento ajuda a comparar lojas, identificar concentra\\u00e7\\u00e3o de volume, localizar varia\\u00e7\\u00f5es bruscas e entender o comportamento do conjunto de acessos cadastrados."),
        image_path=OUT_DIR / "admin-faturamento.png",
        bullets=[
            u("Compare volume bruto, quantidade de vendas e m\\u00e9dia por estabelecimento."),
            u("Use o comparativo para localizar lojas sem movimento ou com queda acentuada."),
            u("Relacione alertas de senha, split e verifica\\u00e7\\u00e3o de e-mail com o desempenho da base."),
            u("A leitura de faturamento serve como triagem para a\\u00e7\\u00f5es operacionais."),
        ],
        note={
            "label": u("Dica"),
            "text": u("quando o volume cair de forma inesperada, revise estabelecimento, acesso, split e poss\\u00edveis pend\\u00eancias cadastrais."),
        },
    )

    doc.add_page_break()
    add_section_page(
        doc,
        u("Perfil"),
        u("Conta do administrador e seguran\\u00e7a"),
        u("A \\u00e1rea de perfil concentra dados pessoais, senha, verifica\\u00e7\\u00e3o de e-mail e identidade visual da conta. Em um ambiente administrativo, isso evita confus\\u00e3o entre acessos e refor\\u00e7a rastreabilidade."),
        image_path=OUT_DIR / "admin-profile.png",
        bullets=[
            u("Revise nome e e-mail sempre que houver troca de respons\\u00e1vel."),
            u("Troque a senha em intervalos regulares ou quando houver suspeita de exposi\\u00e7\\u00e3o."),
            u("Reenvie a verifica\\u00e7\\u00e3o quando o e-mail estiver pendente."),
            u("Mantenha o logotipo e os dados da conta coerentes com a identidade atual."),
        ],
    )

    doc.add_page_break()
    add_simple_page(
        doc,
        u("Opera\\u00e7\\u00e3o di\\u00e1ria"),
        u("Checklist do administrador"),
        [
            u("Uma rotina simples reduz erro de cadastro e evita altera\\u00e7\\u00f5es sem confer\\u00eancia. Comece sempre pelo painel e pelo per\\u00edodo de consulta."),
            u("Depois siga para a base: estabelecimentos, vendedores e regras de split que precisem de ajuste."),
            u("Se houver mudan\\u00e7a sens\\u00edvel, prefira revisar, exportar e validar antes de editar."),
        ],
        bullets=[
            u("Abrir o dashboard e confirmar o per\\u00edodo."),
            u("Verificar transa\\u00e7\\u00f5es, boletos e alertas."),
            u("Revisar estabelecimentos com pend\\u00eancias."),
            u("Criar ou ajustar acessos de vendedores."),
            u("Checar regras de split e exporta\\u00e7\\u00f5es pendentes."),
        ],
        note={
            "label": u("Boa pr\\u00e1tica"),
            "text": u("fa\\u00e7a uma altera\\u00e7\\u00e3o por vez e confirme o efeito na tela antes de continuar para o pr\\u00f3ximo item."),
        },
    )

    doc.add_page_break()
    add_simple_page(
        doc,
        u("Encerramento"),
        u("Resumo final"),
        [
            u("Este manual consolida as opera\\u00e7\\u00f5es que o administrador precisa dominar para manter a base do Checkout Juntter organizada e segura."),
            u("Se a equipe evoluir o sistema, o pr\\u00f3ximo passo natural \\u00e9 atualizar este manual junto com a vers\\u00e3o do super admin para manter consist\\u00eancia entre perfis."),
        ],
        bullets=[
            u("Painel e monitoramento"),
            u("Estabelecimentos e exporta\\u00e7\\u00f5es"),
            u("Split pr\\u00e9 e vendedores"),
            u("Perfil e seguran\\u00e7a"),
        ],
        note={
            "label": u("Pr\\u00f3ximo passo"),
            "text": u("se quiser, eu tamb\\u00e9m posso gerar a vers\\u00e3o do super admin no mesmo padr\\u00e3o e fechar a cole\\u00e7\\u00e3o completa dos manuais."),
        },
    )

    for section in doc.sections:
        section.page_width = PAGE_W
        section.page_height = PAGE_H
        section.left_margin = Mm(MARGIN_MM)
        section.right_margin = Mm(MARGIN_MM)
        section.top_margin = Mm(12)
        section.bottom_margin = Mm(12)

    output = OUT_DIR / "manual-administrador-checkout-juntter.docx"
    doc.save(str(output))
    return output


if __name__ == "__main__":
    build_images()
    print(build_docx())
