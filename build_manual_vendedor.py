from pathlib import Path
import textwrap

from PIL import Image, ImageDraw, ImageFilter, ImageFont, ImageOps
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


def u(value: str) -> str:
    return value.encode("ascii").decode("unicode_escape")


BASE = Path(r"C:\Users\User\Desktop\juntter")
ASSETS = BASE / "docs_unpacked" / "docs"
OUT_DIR = Path(r"C:\xampp\htdocs\juntter\docs")
OUT_DIR.mkdir(parents=True, exist_ok=True)

FONT_REGULAR = r"C:\Windows\Fonts\arial.ttf"
FONT_BOLD = r"C:\Windows\Fonts\arialbd.ttf"

PAGE_W = Mm(210)
PAGE_H = Mm(297)
MARGIN_MM = 17
CONTENT_W = Inches(8.27 - (MARGIN_MM / 25.4) * 2)

CANVAS_W = 1654
CANVAS_H = 2339

GOLD = "FFC400"
BLACK = "111111"
MID = "666666"
LIGHT = "#F7F7F6"


def pcolor(color: str) -> str:
    if color.startswith("#"):
        return color
    return f"#{color}"


def font(path: str, size: int) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(path, size=size)


def add_text(draw: ImageDraw.ImageDraw, xy, text, ft, fill, anchor="la", spacing=8, align="left"):
    draw.multiline_text(
        xy,
        text,
        font=ft,
        fill=pcolor(fill),
        spacing=spacing,
        anchor=anchor,
        align=align,
    )


def rounded_mask(size, radius):
    mask = Image.new("L", size, 0)
    drawer = ImageDraw.Draw(mask)
    drawer.rounded_rectangle((0, 0, size[0] - 1, size[1] - 1), radius=radius, fill=255)
    return mask


def fit_image(img, target_size):
    return ImageOps.fit(img, target_size, method=Image.Resampling.LANCZOS)


def create_cover():
    source = Image.open(ASSETS / "homepage-screenshot.png").convert("RGB")
    hero_crop = source.crop((0, 0, source.width, 1500))

    bg = Image.new("RGB", (CANVAS_W, CANVAS_H), "#F7F7F6")
    overlay = Image.new("RGBA", (CANVAS_W, CANVAS_H), (255, 255, 255, 0))
    overlay_draw = ImageDraw.Draw(overlay)
    overlay_draw.rounded_rectangle((115, 90, CANVAS_W - 115, 1145), radius=34, fill=(255, 255, 255, 255))
    bg = Image.alpha_composite(bg.convert("RGBA"), overlay).convert("RGB")

    hero = fit_image(hero_crop, (CANVAS_W - 290, 920))
    hero_shadow = Image.new("RGBA", (hero.size[0] + 24, hero.size[1] + 24), (0, 0, 0, 0))
    shadow_draw = ImageDraw.Draw(hero_shadow)
    shadow_draw.rounded_rectangle((12, 12, hero.size[0] + 11, hero.size[1] + 11), radius=30, fill=(0, 0, 0, 55))
    hero_shadow = hero_shadow.filter(ImageFilter.GaussianBlur(16))

    bg_rgba = bg.convert("RGBA")
    bg_rgba.alpha_composite(hero_shadow, (98, 80))
    hero_mask = rounded_mask(hero.size, 28)
    bg_rgba.paste(hero.convert("RGBA"), (110, 92), hero_mask)

    draw = ImageDraw.Draw(bg_rgba)
    ft_logo = font(FONT_BOLD, 34)
    ft_kicker = font(FONT_BOLD, 28)
    ft_title = font(FONT_BOLD, 72)
    ft_sub = font(FONT_REGULAR, 34)
    ft_small = font(FONT_REGULAR, 24)
    ft_chip = font(FONT_BOLD, 23)

    logo = Image.open(ASSETS / "logo" / "juntter_png_256.png").convert("RGBA")
    logo = ImageOps.contain(logo, (170, 170), method=Image.Resampling.LANCZOS)
    bg_rgba.alpha_composite(logo, (118, 36))
    add_text(draw, (300, 56), u("Checkout Juntter"), ft_logo, BLACK)

    draw.rounded_rectangle((110, 1125, CANVAS_W - 110, 2210), radius=34, fill="#FFFFFF", outline="#EFEFEF", width=2)
    draw.rectangle((150, 1176, 308, 1182), fill=f"#{GOLD}")
    add_text(draw, (150, 1215), u("MANUAL DO VENDEDOR"), ft_kicker, f"#{GOLD}")
    add_text(draw, (150, 1272), u("Checkout Juntter"), ft_title, BLACK)
    subtitle = textwrap.fill(
        u(
        "Opera\\u00e7\\u00e3o di\\u00e1ria para acessar o sistema, emitir cobran\\u00e7as, cadastrar produtos, "
        "publicar links e acompanhar vendas."
        ),
        width=44,
    )
    add_text(draw, (150, 1388), subtitle, ft_sub, MID)

    chips = [
        (u("Base editorial"), u("Template visual do manual")),
        (u("P\\u00fablico-alvo"), u("Perfil vendedor")),
        (u("Escopo"), u("Cobran\\u00e7a, checkout e perfil")),
    ]
    x = 150
    y = 1548
    for head, body in chips:
        w = 460
        h = 155
        draw.rounded_rectangle((x, y, x + w, y + h), radius=24, fill="#FAFAFA", outline="#E7E7E7", width=2)
        draw.rounded_rectangle((x + 22, y + 22, x + 84, y + 84), radius=18, fill=f"#{GOLD}")
        add_text(draw, (x + 28, y + 31), head[:1], ft_chip, BLACK)
        add_text(draw, (x + 105, y + 28), head, ft_chip, BLACK)
        add_text(draw, (x + 105, y + 68), body, ft_small, MID)
        x += w + 26

    draw.line((150, 1745, CANVAS_W - 150, 1745), fill="#E6E6E6", width=3)
    add_text(draw, (150, 1785), u("Vers\\u00e3o de refer\\u00eancia: julho de 2026"), ft_small, MID)
    add_text(
        draw,
        (150, 1830),
        u("Este manual re\\u00fane as rotinas vis\\u00edveis no sistema para o vendedor operar com seguran\\u00e7a e consist\\u00eancia."),
        ft_small,
        BLACK,
    )

    bg_rgba.convert("RGB").save(OUT_DIR / "manual-cover.png", quality=95)


def create_card_montage(items, outfile, title, subtitle, layout=None):
    if layout is None:
        if len(items) == 1:
            layout = [[0]]
        elif len(items) == 2:
            layout = [[0, 1]]
        elif len(items) == 3:
            layout = [[0], [1, 2]]
        else:
            layout = [[0, 1], [2, 3]]

    rows = len(layout)
    cols = max(len(row) for row in layout)
    canvas_w = 2400 if cols >= 2 else 1800
    canvas_h = 1900 if rows >= 2 else 1320
    if rows == 1 and cols == 1:
        canvas_w = 1800
        canvas_h = 1320

    canvas = Image.new("RGB", (canvas_w, canvas_h), LIGHT)
    draw = ImageDraw.Draw(canvas)
    ft_title = font(FONT_BOLD, 48)
    ft_sub = font(FONT_REGULAR, 26)
    ft_label = font(FONT_BOLD, 24)
    ft_caption = font(FONT_REGULAR, 22)

    add_text(draw, (80, 60), title, ft_title, BLACK)
    add_text(draw, (80, 125), subtitle, ft_sub, MID)

    top = 220
    bottom = 70
    left = 70
    right = 70
    gap_x = 34
    gap_y = 34
    avail_w = canvas_w - left - right
    avail_h = canvas_h - top - bottom
    cell_w = int((avail_w - gap_x * (cols - 1)) / cols)
    cell_h = int((avail_h - gap_y * (rows - 1)) / rows)

    for r, row in enumerate(layout):
        for c, idx in enumerate(row):
            item = items[idx]
            x = left + c * (cell_w + gap_x)
            y = top + r * (cell_h + gap_y)
            card = Image.new("RGBA", (cell_w, cell_h), "white")
            card_draw = ImageDraw.Draw(card)
            card_draw.rounded_rectangle((0, 0, cell_w - 1, cell_h - 1), radius=28, fill="white", outline="#E6E6E6", width=2)
            shot = Image.open(item["path"]).convert("RGB")
            img_h = cell_h - 96
            img_w = cell_w - 34
            fitted = fit_image(shot, (img_w, img_h))
            card.paste(fitted, (17, 16))
            card_draw.rectangle((17, img_h + 26, cell_w - 17, img_h + 28), fill=f"#{GOLD}")
            add_text(card_draw, (18, img_h + 38), f"{item['num']}. {item['label']}", ft_label, BLACK)
            add_text(card_draw, (18, img_h + 68), item.get("caption", ""), ft_caption, MID)
            canvas.paste(card.convert("RGB"), (x, y))

    canvas.save(outfile, quality=95)


def create_single_image_page(item, outfile, title, subtitle):
    canvas = Image.new("RGB", (1800, 1320), LIGHT)
    draw = ImageDraw.Draw(canvas)
    ft_title = font(FONT_BOLD, 48)
    ft_sub = font(FONT_REGULAR, 26)
    ft_label = font(FONT_BOLD, 24)
    ft_caption = font(FONT_REGULAR, 22)
    add_text(draw, (80, 60), title, ft_title, BLACK)
    add_text(draw, (80, 125), subtitle, ft_sub, MID)

    x, y, cell_w, cell_h = 70, 220, 1660, 980
    card = Image.new("RGBA", (cell_w, cell_h), "white")
    card_draw = ImageDraw.Draw(card)
    card_draw.rounded_rectangle((0, 0, cell_w - 1, cell_h - 1), radius=28, fill="white", outline="#E6E6E6", width=2)
    shot = Image.open(item["path"]).convert("RGB")
    fitted = fit_image(shot, (cell_w - 34, cell_h - 120))
    card.paste(fitted, (17, 16))
    card_draw.rectangle((17, cell_h - 86, cell_w - 17, cell_h - 84), fill=f"#{GOLD}")
    add_text(card_draw, (18, cell_h - 76), f"{item['num']}. {item['label']}", ft_label, BLACK)
    add_text(card_draw, (18, cell_h - 46), item.get("caption", ""), ft_caption, MID)
    canvas.paste(card.convert("RGB"), (x, y))
    canvas.save(outfile, quality=95)


def create_two_image_page(items, outfile, title, subtitle):
    create_card_montage(items, outfile, title, subtitle, layout=[[0, 1]])


def create_three_image_page(items, outfile, title, subtitle):
    create_card_montage(items, outfile, title, subtitle, layout=[[0], [1, 2]])


def set_run_font(run, name="Arial", size=11, bold=False, color=BLACK, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:cs"), name)


def set_paragraph_spacing(paragraph, before=0, after=0, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_header_and_footer(section):
    header = section.header
    h = header.paragraphs[0]
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_run = h.add_run()
    logo_run.add_picture(str(ASSETS / "logo" / "juntter_png_256.png"), width=Inches(0.45))
    title_run = h.add_run(u("  Manual do Vendedor | Checkout Juntter"))
    set_run_font(title_run, size=9, bold=True, color=BLACK)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_text = fp.add_run(u("Checkout Juntter - Manual do Vendedor  "))
    set_run_font(footer_text, size=8.5, color=MID)
    add_page_number(fp)
    for run in fp.runs:
        if run.text == "PAGE":
            set_run_font(run, size=8.5, color=MID)


def style_doc(doc):
    styles = doc.styles
    for style_name in ["Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
        try:
            style = styles[style_name]
            style.font.name = "Arial"
            style.font.size = Pt(11)
            style.font.color.rgb = RGBColor.from_string(BLACK)
        except Exception:
            pass

    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(5)

    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor.from_string(BLACK)

    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = RGBColor.from_string(BLACK)

    styles["Heading 3"].font.name = "Arial"
    styles["Heading 3"].font.size = Pt(11)
    styles["Heading 3"].font.bold = True


def add_cover_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(str(OUT_DIR / "manual-cover.png"), width=CONTENT_W)


def add_intro_page(doc):
    p = doc.add_paragraph()
    r = p.add_run(u("DIRE\\u00c7\\u00c3O EDITORIAL"))
    set_run_font(r, size=9, bold=True, color=GOLD)

    p = doc.add_paragraph()
    r = p.add_run(u("Como ler este manual"))
    set_run_font(r, size=22, bold=True, color=BLACK)

    for text in [
        u("O manual foi escrito para apoiar a rotina real do vendedor. Cada se\\u00e7\\u00e3o responde a tr\\u00eas perguntas: o que fazer, quando usar e como concluir a tarefa."),
        u("A linguagem \\u00e9 direta, com passos numerados quando existe sequ\\u00eancia operacional e com observa\\u00e7\\u00f5es curtas quando h\\u00e1 risco, depend\\u00eancia ou cuidado especial."),
        u("O padr\\u00e3o visual privilegia preto, branco e dourado, com blocos limpos e poucas distra\\u00e7\\u00f5es, para manter a leitura r\\u00e1pida no dia a dia."),
    ]:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=6, line=1.25)
        r = p.add_run(text)
        set_run_font(r, size=10.5, color="444444")

    for text in [
        u("Comece pela p\\u00e1gina inicial e confira o per\\u00edodo exibido nos indicadores."),
        u("Use Cobran\\u00e7a para acompanhar pagamentos e movimenta\\u00e7\\u00f5es financeiras."),
        u("Use Checkout para cadastrar produtos, publicar links e acompanhar vendas."),
        u("Use Perfil para trocar senha, reenviar verifica\\u00e7\\u00e3o e atualizar o logotipo."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        set_paragraph_spacing(p, after=1, line=1.2)
        p.paragraph_format.left_indent = Inches(0.18)
        r = p.add_run(text)
        set_run_font(r, size=10.5, color=BLACK)

    p = doc.add_paragraph()
    set_shading(p, "FFF7D1")
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(u("Princ\\u00edpio central: "))
    set_run_font(r, size=10.5, bold=True, color=BLACK)
    r = p.add_run(u("se uma tela pede aten\\u00e7\\u00e3o, ela deve trazer a a\\u00e7\\u00e3o principal, o resultado esperado e o pr\\u00f3ximo passo em poucos segundos."))
    set_run_font(r, size=10.5, color=BLACK)


def add_simple_page(doc, kicker, title, paragraphs, bullets=None, note=None):
    p = doc.add_paragraph()
    r = p.add_run(kicker.upper())
    set_run_font(r, size=9, bold=True, color=GOLD)

    p = doc.add_paragraph()
    r = p.add_run(title)
    set_run_font(r, size=22, bold=True, color=BLACK)

    for text in paragraphs:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=6, line=1.25)
        r = p.add_run(text)
        set_run_font(r, size=10.5, color="444444")

    if bullets:
        for text in bullets:
            p = doc.add_paragraph(style="List Bullet")
            set_paragraph_spacing(p, after=1, line=1.2)
            r = p.add_run(text)
            set_run_font(r, size=10.5, color=BLACK)

    if note:
        p = doc.add_paragraph()
        set_shading(p, "F0F0F0")
        r = p.add_run(note["label"] + ": ")
        set_run_font(r, size=10.5, bold=True, color=BLACK)
        r = p.add_run(note["text"])
        set_run_font(r, size=10.5, color=BLACK)


def add_section_page(doc, kicker, title, intro, image_path=None, bullets=None, steps=None, callout=None):
    p = doc.add_paragraph()
    r = p.add_run(kicker.upper())
    set_run_font(r, size=9, bold=True, color=GOLD)

    p = doc.add_paragraph()
    r = p.add_run(title)
    set_run_font(r, size=22, bold=True, color=BLACK)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=8, line=1.25)
    r = p.add_run(intro)
    set_run_font(r, size=10.5, color="444444")

    if image_path:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(5)
        p.add_run().add_picture(str(image_path), width=CONTENT_W)

    if callout:
        p = doc.add_paragraph()
        set_shading(p, "FFF7D1")
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(callout["label"] + ": ")
        set_run_font(r, size=10.5, bold=True, color=BLACK)
        r = p.add_run(callout["text"])
        set_run_font(r, size=10.5, color=BLACK)

    if steps:
        for text in steps:
            p = doc.add_paragraph(style="List Number")
            set_paragraph_spacing(p, after=2, line=1.2)
            p.paragraph_format.left_indent = Inches(0.25)
            r = p.add_run(text)
            set_run_font(r, size=10.5, color=BLACK)

    if bullets:
        for text in bullets:
            p = doc.add_paragraph(style="List Bullet")
            set_paragraph_spacing(p, after=1, line=1.2)
            p.paragraph_format.left_indent = Inches(0.18)
            r = p.add_run(text)
            set_run_font(r, size=10.5, color=BLACK)


def build_images():
    create_cover()

    create_three_image_page(
        [
            {
                "path": ASSETS / "screenshots" / "login.png",
                "num": "01",
                "label": u("Acesso"),
                "caption": u("Entrar com e-mail e senha."),
            },
            {
                "path": ASSETS / "screenshots" / "dasboard-opened.png",
                "num": "02",
                "label": u("Painel aberto"),
                "caption": u("Indicadores e filtros do per\\u00edodo."),
            },
            {
                "path": ASSETS / "screenshots" / "dashboard-closed.png",
                "num": "03",
                "label": u("Painel fechado"),
                "caption": u("Modo compacto para leitura rapida."),
            },
        ],
        OUT_DIR / "page-access-dashboard.png",
        u("Primeiro acesso e painel"),
        u("Acesso, verifica\\u00e7\\u00e3o de e-mail e leitura dos indicadores do vendedor."),
    )

    create_card_montage(
        [
            {
                "path": ASSETS / "screenshots" / "history.png",
                "num": "01",
                "label": u("Hist\\u00f3rico"),
                "caption": u("Filtre cobran\\u00e7as por per\\u00edodo e status."),
            },
            {
                "path": ASSETS / "screenshots" / "pix-create.png",
                "num": "02",
                "label": u("PIX"),
                "caption": u("Cria\\u00e7\\u00e3o e acompanhamento da cobran\\u00e7a instant\\u00e2nea."),
            },
            {
                "path": ASSETS / "screenshots" / "credit-card-charge-creation-form.png",
                "num": "03",
                "label": u("Cartao"),
                "caption": u("Entrada de dados e valor da cobran\\u00e7a."),
            },
            {
                "path": ASSETS / "screenshots" / "boleto-transaction-creation-form.png",
                "num": "04",
                "label": u("Boleto"),
                "caption": u("Emissao e consulta do boleto."),
            },
        ],
        OUT_DIR / "page-charges.png",
        u("Cobrancas"),
        u("Hist\\u00f3rico, PIX, cart\\u00e3o de cr\\u00e9dito e boleto em uma vis\\u00e3o operacional \\u00fanica."),
    )

    create_three_image_page(
        [
            {
                "path": ASSETS / "screenshots" / "balance.png",
                "num": "01",
                "label": u("Saldo e extrato"),
                "caption": u("Disponivel, em transito e em processamento."),
            },
            {
                "path": ASSETS / "screenshots" / "simulate-transaction.png",
                "num": "02",
                "label": u("Simula\\u00e7\\u00e3o"),
                "caption": u("Teste fluxos antes da opera\\u00e7\\u00e3o real."),
            },
            {
                "path": ASSETS / "screenshots" / "pix-generate-qrcode.png",
                "num": "03",
                "label": u("PIX"),
                "caption": u("Refer\\u00eancia visual para o fluxo PIX."),
            },
        ],
        OUT_DIR / "page-financial.png",
        u("Saldo, extrato e simulacao"),
        u("Consulta financeira, confer\\u00eancia de movimentos e testes de cobran\\u00e7a."),
    )

    create_single_image_page(
        {
            "path": ASSETS / "screenshots" / "checkout-products.png",
            "num": "01",
            "label": u("Produtos"),
            "caption": u("Cadastre, edite e mantenha o catalogo vinculado ao checkout."),
        },
        OUT_DIR / "page-products.png",
        u("Produtos"),
        u("Cadastro e manutencao dos itens usados nos links de checkout."),
    )

    create_two_image_page(
        [
            {
                "path": ASSETS / "screenshots" / "checkout-links-list.png",
                "num": "01",
                "label": u("Lista de links"),
                "caption": u("Veja status, preco, vendas e acoes rapidas."),
            },
            {
                "path": ASSETS / "screenshots" / "checkout-links-creation-form.png",
                "num": "02",
                "label": u("Criacao"),
                "caption": u("Produto, quantidade, preco e meios de pagamento."),
            },
        ],
        OUT_DIR / "page-checkout-links-1.png",
        u("Links de checkout"),
        u("Publica\\u00e7\\u00e3o do endere\\u00e7o de venda e configura\\u00e7\\u00e3o comercial b\\u00e1sica."),
    )

    create_two_image_page(
        [
            {
                "path": ASSETS / "screenshots" / "checkout-links-creation-form2.png",
                "num": "03",
                "label": u("Apar\\u00eancia"),
                "caption": u("Cores, mensagens e URLs de retorno."),
            },
            {
                "path": ASSETS / "screenshots" / "checkout-link-payment-wizard.png",
                "num": "04",
                "label": u("Fluxo do cliente"),
                "caption": u("Visao da jornada ate o pagamento."),
            },
        ],
        OUT_DIR / "page-checkout-links-2.png",
        u("Links de checkout - continuacao"),
        u("Ajustes visuais e leitura do fluxo de pagamento do cliente."),
    )

    create_two_image_page(
        [
            {
                "path": ASSETS / "screenshots" / "credit-card-transaction-list.png",
                "num": "01",
                "label": u("Vendas em cart\\u00e3o"),
                "caption": u("Lista e leitura dos pagamentos conclu\\u00eddos."),
            },
            {
                "path": ASSETS / "screenshots" / "boleto-transaction-list.png",
                "num": "02",
                "label": u("Vendas em boleto"),
                "caption": u("Acompanhamento dos boletos gerados."),
            },
        ],
        OUT_DIR / "page-sales.png",
        u("Vendas do link"),
        u("Acompanhamento de pedidos e leitura do resultado financeiro."),
    )

    create_single_image_page(
        {
            "path": ASSETS / "screenshots" / "profile.png",
            "num": "01",
            "label": u("Perfil"),
            "caption": u("Dados pessoais, senha, logotipo e verifica\\u00e7\\u00e3o de e-mail."),
        },
        OUT_DIR / "page-profile.png",
        u("Perfil e seguran\\u00e7a"),
        u("Manuten\\u00e7\\u00e3o da conta, autentica\\u00e7\\u00e3o e identidade visual do vendedor."),
    )


def build_docx():
    doc = Document()
    style_doc(doc)

    sec = doc.sections[0]
    sec.page_width = PAGE_W
    sec.page_height = PAGE_H
    sec.left_margin = Mm(MARGIN_MM)
    sec.right_margin = Mm(MARGIN_MM)
    sec.top_margin = Mm(12)
    sec.bottom_margin = Mm(12)
    sec.different_first_page_header_footer = True
    add_header_and_footer(sec)

    add_cover_page(doc)
    doc.add_page_break()

    add_intro_page(doc)
    doc.add_page_break()

    add_section_page(
        doc,
        u("Primeiro acesso"),
        u("Acesso ao sistema e leitura do painel"),
        u("Entre com seu e-mail e senha, conclua a verifica\\u00e7\\u00e3o de e-mail quando o sistema solicitar e siga para o painel inicial. A primeira leitura deve ser sempre feita pelo per\\u00edodo exibido no topo da tela."),
        image_path=OUT_DIR / "page-access-dashboard.png",
        steps=[
            u("Acesse /app/login e informe suas credenciais."),
            u("Se houver pend\\u00eancia de verifica\\u00e7\\u00e3o, conclua o processo antes de operar."),
            u("Quando a senha precisar ser trocada, atualize-a antes de prosseguir."),
            u("No painel, confirme mes, ano, saldo e volume de transacoes antes de executar qualquer acao."),
        ],
        callout={"label": u("Dica"), "text": u("se a conta ainda n\\u00e3o estiver verificada, fa\\u00e7a esse ajuste antes de emitir cobran\\u00e7as ou publicar links.")},
    )

    doc.add_page_break()
    add_section_page(
        doc,
        u("Cobranca"),
        u("Cobran\\u00e7as: hist\\u00f3rico, PIX, cart\\u00e3o e boleto"),
        u("A \\u00e1rea de cobran\\u00e7a re\\u00fane as opera\\u00e7\\u00f5es financeiras que o vendedor precisa acompanhar no dia a dia. Use o hist\\u00f3rico para revisar transa\\u00e7\\u00f5es e as telas espec\\u00edficas para criar novas cobran\\u00e7as por meio de pagamento."),
        image_path=OUT_DIR / "page-charges.png",
        bullets=[
            u("No hist\\u00f3rico, filtre por per\\u00edodo, pesquise pelo cliente e abra o detalhe da transa\\u00e7\\u00e3o quando precisar validar uma informa\\u00e7\\u00e3o."),
            u("No PIX, informe o valor e acompanhe a gera\\u00e7\\u00e3o da cobran\\u00e7a at\\u00e9 a confirma\\u00e7\\u00e3o do pagamento."),
            u("No cart\\u00e3o de cr\\u00e9dito, revise os dados e o valor antes de submeter a cobran\\u00e7a."),
            u("No boleto, confira vencimento, instru\\u00e7\\u00f5es e status de pagamento antes de encaminhar ao cliente."),
        ],
        callout={"label": u("Importante"), "text": u("uma cobran\\u00e7a confirmada pode gerar movimenta\\u00e7\\u00f5es financeiras reais. Revise valor, cliente e meio de pagamento antes de concluir.")},
    )

    doc.add_page_break()
    add_section_page(
        doc,
        u("Saldo e testes"),
        u("Saldo, extrato, PIX de saida e simulacao"),
        u("Use essa parte da plataforma para conferir disponibilidade financeira, analisar lancamentos e testar o comportamento das transacoes antes de operar em producao."),
        image_path=OUT_DIR / "page-financial.png",
        steps=[
            u("Abra o saldo e extrato para verificar disponivel, em transito e em processamento."),
            u("Aplique o filtro de per\\u00edodo quando quiser revisar a movimenta\\u00e7\\u00e3o de um m\\u00eas espec\\u00edfico."),
            u("Use a simula\\u00e7\\u00e3o para validar o comportamento esperado antes de gerar uma cobran\\u00e7a real."),
            u("No PIX de sa\\u00edda, valide dados do destinat\\u00e1rio e confirme a opera\\u00e7\\u00e3o apenas depois da revis\\u00e3o final."),
        ],
        callout={"label": u("Aten\\u00e7\\u00e3o"), "text": u("dependendo da configura\\u00e7\\u00e3o da conta, a sa\\u00edda via PIX pode exigir assinatura eletr\\u00f4nica e confirma\\u00e7\\u00e3o adicional.")},
    )

    doc.add_page_break()
    add_section_page(
        doc,
        u("Produtos"),
        u("Cadastro e manutencao de produtos"),
        u("Os produtos alimentam os links de checkout. Mantenha nomes, precos, imagens e status coerentes com o catalogo que o cliente vera na hora da compra."),
        image_path=OUT_DIR / "page-products.png",
        steps=[
            u("Clique em Novo produto quando precisar incluir um novo item no catalogo."),
            u("Preencha nome, resumo, descricao, preco e SKU quando existir uma referencia interna."),
            u("Envie uma imagem quando ela for necessaria para a apresentacao comercial."),
            u("Altere o status para ativo ou inativo conforme o produto esteja ou n\\u00e3o dispon\\u00edvel para venda."),
            u("Exclua apenas quando o item n\\u00e3o for mais necess\\u00e1rio no cat\\u00e1logo."),
        ],
        callout={"label": u("Exportacao"), "text": u("a base de clientes pode ser exportada em .xls a partir do painel do vendedor, reunindo documento, transacoes, valor total e datas da primeira e da ultima compra.")},
    )

    doc.add_page_break()
    add_section_page(
        doc,
        u("Checkout"),
        u("Criacao de links de checkout"),
        u("Os links de checkout sao a principal forma de publicar a oferta do vendedor. O fluxo comeca pela selecao do produto e termina com o link publico pronto para ser compartilhado."),
        image_path=OUT_DIR / "page-checkout-links-1.png",
        steps=[
            u("Abra a listagem e clique em Novo link."),
            u("Escolha o produto que vai sustentar a oferta."),
            u("Defina quantidade, preco unitario e status do link."),
            u("Selecione os meios de pagamento permitidos e aplique descontos, se necessario."),
            u("Configure frete, expira\\u00e7\\u00e3o e dados pr\\u00e9-preenchidos quando a opera\\u00e7\\u00e3o exigir."),
        ],
        callout={"label": u("Dica"), "text": u("copie e teste o endere\\u00e7o p\\u00fablico antes de enviar ao cliente para evitar publica\\u00e7\\u00e3o com configura\\u00e7\\u00e3o incompleta.")},
    )

    doc.add_page_break()
    add_section_page(
        doc,
        u("Checkout"),
        u("Ajustes visuais e publicacao do link"),
        u("Na segunda etapa, o vendedor refina a aparencia do checkout e decide como o cliente deve ser conduzido ate o pagamento."),
        image_path=OUT_DIR / "page-checkout-links-2.png",
        bullets=[
            u("Use o tema visual para alinhar o checkout com a identidade da loja."),
            u("Ajuste cores, mensagens e URLs de retorno conforme a campanha."),
            u("Revise o status do link e mantenha-o ativo apenas quando ele estiver pronto para receber trafego."),
            u("Quando mudar o produto ou a oferta, atualize a configura\\u00e7\\u00e3o e recopie o endere\\u00e7o p\\u00fablico."),
        ],
        callout={"label": u("Disponibilidade"), "text": u("um link pode aparecer indisponivel se estiver desativado, expirado, associado a um produto inativo ou fora do perfil esperado do vendedor.")},
    )

    doc.add_page_break()
    add_section_page(
        doc,
        u("Vendas"),
        u("Acompanhamento das vendas do link"),
        u("Depois da publicacao, a rotina passa a ser acompanhamento. A tela de vendas mostra o desempenho do link e permite abrir o detalhe de cada pedido para revisao completa."),
        image_path=OUT_DIR / "page-sales.png",
        bullets=[
            u("Na listagem, confira quantidade de pedidos, valor vendido e status das transacoes."),
            u("No detalhe, valide cliente, endereco de entrega, meio de pagamento e dados do pedido."),
            u("Use a leitura de vendas para cruzar o que foi publicado com o que foi efetivamente pago."),
            u("Se houver diverg\\u00eancia, volte \\u00e0 configura\\u00e7\\u00e3o do produto, do link ou da cobran\\u00e7a."),
        ],
        callout={"label": u("Resultado esperado"), "text": u("o vendedor consegue ligar a oferta publicada ao pedido concluido e ao valor que entra no fluxo financeiro.")},
    )

    doc.add_page_break()
    add_section_page(
        doc,
        u("Perfil"),
        u("Perfil da conta, senha e logotipo"),
        u("A \\u00e1rea de perfil concentra os ajustes pessoais e de seguran\\u00e7a da conta. \\u00c9 o lugar para manter dados atualizados e garantir que a identidade visual da empresa esteja correta."),
        image_path=OUT_DIR / "page-profile.png",
        steps=[
            u("Revise nome, e-mail e status de verifica\\u00e7\\u00e3o sempre que houver mudan\\u00e7a na conta."),
            u("Use a troca de senha quando a pol\\u00edtica de acesso exigir atualiza\\u00e7\\u00e3o."),
            u("Reenvie o e-mail de verifica\\u00e7\\u00e3o caso a conta ainda esteja pendente."),
            u("Envie ou remova o logotipo da empresa conforme a marca atual."),
        ],
        callout={"label": u("Observacao"), "text": u("o sistema aceita JPG, PNG e WEBP para o logotipo, com melhor resultado quando a imagem ja esta preparada para uso em tela.")},
    )

    doc.add_page_break()
    add_simple_page(
        doc,
        u("Opera\\u00e7\\u00e3o di\\u00e1ria"),
        u("Checklist de uso do vendedor"),
        [
            u("Uma rotina simples evita erros e ajuda a manter a opera\\u00e7\\u00e3o previs\\u00edvel. Antes de iniciar o dia, confira o painel, valide o per\\u00edodo e confirme o saldo dispon\\u00edvel."),
            u("Depois da leitura inicial, siga para a tarefa do momento: emitir cobran\\u00e7a, publicar link, revisar vendas ou ajustar o cat\\u00e1logo."),
            u("Ao final da execu\\u00e7\\u00e3o, revise o status do item criado, copie o endere\\u00e7o p\\u00fablico quando houver publica\\u00e7\\u00e3o e confirme se o cliente recebeu a informa\\u00e7\\u00e3o correta."),
        ],
        bullets=[
            u("Entrar no sistema e conferir o per\\u00edodo do painel."),
            u("Revisar saldos, transa\\u00e7\\u00f5es e cobran\\u00e7as recentes."),
            u("Atualizar produtos quando o cat\\u00e1logo mudar."),
            u("Publicar ou revisar links de checkout antes de compartilhar."),
            u("Manter perfil, senha e logotipo sempre atualizados."),
        ],
        note={"label": u("Boas pr\\u00e1ticas"), "text": u("n\\u00e3o deixe links expirados em circula\\u00e7\\u00e3o, n\\u00e3o publique produtos inativos e sempre valide os dados do cliente antes de concluir uma cobran\\u00e7a.")},
    )

    doc.add_page_break()
    add_simple_page(
        doc,
        u("Encerramento"),
        u("Resumo final"),
        [
            u("Este manual foi organizado para cobrir o conjunto de opera\\u00e7\\u00f5es que o vendedor precisa conhecer para operar o Checkout Juntter com seguran\\u00e7a."),
            u("Se alguma tela do sistema sofrer atualiza\\u00e7\\u00e3o visual, preserve a mesma ordem operacional: acessar, conferir, configurar, publicar, acompanhar e ajustar."),
            u("Quando a equipe precisar de outra vers\\u00e3o do manual, o pr\\u00f3ximo passo natural \\u00e9 criar as edi\\u00e7\\u00f5es do administrador e do super admin com o mesmo padr\\u00e3o editorial."),
        ],
        bullets=[
            u("Acesso e painel"),
            u("Cobran\\u00e7as e saldo"),
            u("Produtos e checkout"),
            u("Vendas e perfil"),
        ],
        note={"label": u("Pr\\u00f3ximo passo"), "text": u("se quiser, eu tamb\\u00e9m posso gerar a vers\\u00e3o do administrador e depois consolidar os tr\\u00eas manuais em uma cole\\u00e7\\u00e3o \\u00fanica.")},
    )

    for section in doc.sections:
        section.page_width = PAGE_W
        section.page_height = PAGE_H
        section.left_margin = Mm(MARGIN_MM)
        section.right_margin = Mm(MARGIN_MM)
        section.top_margin = Mm(12)
        section.bottom_margin = Mm(12)

    output_path = OUT_DIR / "manual-vendedor-checkout-juntter.docx"
    doc.save(str(output_path))
    return output_path


if __name__ == "__main__":
    build_images()
    path = build_docx()
    print(path)
